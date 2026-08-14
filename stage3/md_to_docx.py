# md_to_docx.py — 论文 md → Word（格式排版控制）
# 处理：标题层级/表格/图片/代码块/引用块/粗体/列表——中文字体（正文宋体/标题黑体）
import re
import sys
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
    'C:/Users/bai/Desktop/AB系统论文储备/多层级语义指纹系统：意图保持的测量—干预—边界闭环（人机判别与推理期干预实证）.md')
OUT = SRC.with_suffix('.docx')
BASE = SRC.parent

# ---------- 样式 ----------
def set_font(run, name_cn='宋体', name_en='Times New Roman', size=10.5, bold=False, color=None, italic=False):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, style=None, size=10.5, bold=False, align=None, space_after=6, indent=False,
             color=None, name_cn='宋体'):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.25
    if indent:
        pf.first_line_indent = Pt(21)  # 首行缩进 2 字符
    # 处理 **粗体** 内联
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_font(r, name_cn=name_cn, size=size, bold=True, color=color)
        else:
            r = p.add_run(part)
            set_font(r, name_cn=name_cn, size=size, bold=bold, color=color)
    return p

def add_heading(doc, text, level):
    # 黑体标题——自定义（不用默认 Heading 的英文样式）
    sizes = {1: 16, 2: 14, 3: 12.5, 4: 11.5}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level <= 2 else 8)
    pf.space_after = Pt(8 if level <= 2 else 5)
    pf.keep_with_next = True
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith('**') and part.endswith('**')
        t = part[2:-2] if bold else part
        r = p.add_run(t)
        set_font(r, name_cn='黑体', size=sizes.get(level, 11), bold=True)
    return p

def add_table(doc, header, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_font(r, name_cn='黑体', size=9, bold=True)
    n_cols = len(header)
    n_rows = len(rows) + 1
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(header):
        cell = tbl.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, name_cn='黑体', size=8.5, bold=True)
        # 表头底色
        shd = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): 'E8EEF7'})
        cell._tc.get_or_add_tcPr().append(shd)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            parts = re.split(r'(\*\*[^*]+\*\*)', str(val))
            for part in parts:
                if not part:
                    continue
                bold = part.startswith('**') and part.endswith('**')
                r = p.add_run(part[2:-2] if bold else part)
                set_font(r, size=8, bold=bold)
    # 列宽：总宽 16cm 均分（超宽表压缩）
    total = sum(len(str(h)) for h in header)
    for j in range(n_cols):
        w = min(max(1.2, 16.0 / n_cols), 6.0)
        for i in range(n_rows):
            tbl.cell(i, j).width = Cm(w)
    return tbl

def add_image(doc, path, caption=None):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(path), width=Cm(15))
    except Exception as e:
        print(f'  图片失败 {path.name}: {str(e)[:60]}')
        return
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_font(r, size=9, color=(0x55, 0x55, 0x55))

# ---------- 解析 ----------
def parse_table_block(lines, i):
    """lines[i] 为表头行——返回 (header, rows, 下一行号)"""
    header = [c.strip().replace('**', '') for c in lines[i].strip().strip('|').split('|')]
    j = i + 1
    while j < len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[j]) and '-' in lines[j]:
        j += 1  # 分隔行
    rows = []
    while j < len(lines) and lines[j].strip().startswith('|'):
        cells = [c.strip() for c in lines[j].strip().strip('|').split('|')]
        while len(cells) < len(header):
            cells.append('')
        rows.append(cells[:len(header)])
        j += 1
    return header, rows, j

def main():
    text = SRC.read_text(encoding='utf-8')
    lines = text.split('\n')
    doc = Document()
    # 页面设置 A4
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    i = 0
    n_table = 0
    n_img = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            in_code = not in_code
            if not in_code:
                i += 1
                continue
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            set_font(r, name_cn='宋体', name_en='Consolas', size=8.5, color=(0x33, 0x33, 0x33))
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # 表格
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i + 1]) and '-' in lines[i + 1]:
            header, rows, j = parse_table_block(lines, i)
            add_table(doc, header, rows)
            n_table += 1
            i = j
            continue

        # 图片
        m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if m:
            img_path = BASE / m.group(2).strip()
            add_image(doc, img_path, m.group(1) or None)
            n_img += 1
            i += 1
            continue

        # 引用块
        if stripped.startswith('>'):
            content = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(5)
            parts = re.split(r'(\*\*[^*]+\*\*)', content)
            for part in parts:
                if not part:
                    continue
                bold = part.startswith('**') and part.endswith('**')
                r = p.add_run(part[2:-2] if bold else part)
                set_font(r, size=9.5, italic=True, color=(0x44, 0x44, 0x44), bold=bold)
            i += 1
            continue

        # 列表
        m = re.match(r'^\s*[-*]\s+(.*)$', stripped)
        if m:
            add_para(doc, m.group(1), size=10, space_after=3, indent=False)
            i += 1
            continue
        m = re.match(r'^\s*\d+\.\s+(.*)$', stripped)
        if m:
            add_para(doc, m.group(1), size=10, space_after=3, indent=False)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线
        if stripped in ('---', '***'):
            i += 1
            continue

        # 正文
        add_para(doc, stripped, size=10.5, space_after=6, indent=True)
        i += 1

    doc.save(OUT)
    print(f'转换完成 ✓ {OUT.name}')
    print(f'  表格 {n_table} 个——图片 {n_img} 张——段落总计 {len(doc.paragraphs)}')

if __name__ == '__main__':
    main()
